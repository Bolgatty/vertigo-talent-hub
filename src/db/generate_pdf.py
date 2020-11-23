from docx2pdf import convert
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from src.db.jd_manager import JDManager as jd
from src.gui.jobs_table_display_gui import JobsTable as jt
from src.tools.thread_pool_jdPDF import ThreadPooljdPDF as tpdf
from tkinter import messagebox as mb
import tkinter as tk

class JobGeneratepdf:

    def generate_pdf(self, root):
        """
            method called when Generate Job button is clicked from job_management initial screen
        """
        if jt.selected_job_id:
            if mb.askyesno("PDF Generate Confirmation",
                           "Are you sure you want to generate pdf for the selected job ?", parent=root):
                for i in jt.issue_key_list:
                    self.pdf_generator(i)
                mb.showinfo("Info", "PDF and DOCx files generated successfully!", parent=root)
        else:
            mb.showerror("Error", "Please select the job you want to generate pdf for", parent=root)

    def pdf_generator(self, issue_key):
        final = Document()
        final.add_picture('src/gui/images/logo.PNG', width=Inches(1.50))
        last_pic = final.paragraphs[-1]
        last_pic.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        final.add_heading('About US', 0)
        about_us = Document('src/gui/about.docx')

        # for element in about_us.element.body:
        #     final.element.body.append(element)

        all_paras = about_us.paragraphs
        print(len(all_paras))

        for para in all_paras:
            print(para.text)
            paragraph = final.add_paragraph(para.text)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(2)

        root = tk.Toplevel()

        #Thread program which inturn calls Job Manager method get_particular_desc, which runs in a separate thread.
        job_dict = tpdf(root, issue_key).workerThread1(issue_key)

        final.add_heading(job_dict['job_id']+' : ' + job_dict['job_title'], 1)

        final.add_heading('Job Description:', 1)
        final.add_paragraph(job_dict['job_description'])

        filename = job_dict['job_id']
        final.save(filename+'.docx')

        convert(filename + '.docx',
                'C:/Users/user-name/vertigo-talent-hunt/give the path where you needed' + filename + '.pdf')

        for line in final.paragraphs:
            with open(filename + 'txt', 'a') as wf:
                wf.write(line.text + "\n")


